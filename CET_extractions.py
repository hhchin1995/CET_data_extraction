from docx.document import Document
from docx import Document
from author import Authors, CorrespondingAuthor
from manuscript import ManuscriptInfo
from typing import List, Dict

#specific to extracting information from word documents
import os
import zipfile
#other tools useful in extracting the information from our document
import re
#to pretty print our xml:
import xml.dom.minidom
import pandas as pd
import openpyxl

from tkinter import Tk
from tkinter.filedialog import askdirectory
from flask import Flask, jsonify, request, render_template, send_from_directory
import uuid
import shutil
import io


class CETExtraction():
    authors: Authors
    corresponding_author: Authors
    manuscript_info: ManuscriptInfo
    paper_id : str
    affiliations: Dict[any, any]
    email: str
    _is_contain_superscripts: bool

    def __init__(self, filename: str):
        document = Document(filename)
        self.paper_id = self._get_paper_id(filename = filename)
        print("\n" + self.paper_id)
        for p, paragraph in enumerate(document.paragraphs):
            if paragraph.style.name == 'CET Authors' or p == 1:
                self.authors, corresponding_author_affiliation_label = self._get_authors_names_2(paragraph)
                break

        for p, paragraph in enumerate(document.paragraphs):
            if p == 0:
                self.manuscript_info = self._get_manuscript_info(filename, paragraph)
                break
        
        self.affiliations = self._get_affiliations(document.paragraphs)
        self.email = self._get_email(document.paragraphs)

        corresponding_author_affiliations = ''
        if corresponding_author_affiliation_label:
            corresponding_author_affiliation_label.reverse()
            for label in corresponding_author_affiliation_label:
                corresponding_author_affiliations += ''.join(self.affiliations[label]) if len(corresponding_author_affiliations) == 0 else "\n" + ''.join(self.affiliations[label])
        else:
            corresponding_author_affiliations += ''.join(self.affiliations)
        
        self.authors.corresponding_author.affiliation = corresponding_author_affiliations
        self.authors.corresponding_author.email = self.email
        
                
    def _get_paper_id(self, filename: str):
        # return filename.split('-')[-1].split('_')[0] + '.pdf' # SCE
        return filename.filename.split('_')[1].split('_')[0] + '.pdf'

    def _get_manuscript_info(self, filename, paragraph):
        page_number = int(self._get_page_number(filename = filename))
        paper_title = paragraph.text
        return ManuscriptInfo(page_no = page_number, paper_title = paper_title)

    def _get_page_number(self, filename):
        try:
            document = zipfile.ZipFile(filename)
            dxml = document.read('docProps/app.xml')
            uglyXml = xml.dom.minidom.parseString(dxml)
            page = uglyXml.getElementsByTagName('Pages')[0].childNodes[0].nodeValue
        except Exception as e:
            # raise
            page = 0
            pass
        return page
    
    def _get_affiliations(self, paragraphs):
        affiliations = {}
        superscript = ''
        for p, paragraph in enumerate(paragraphs):
            if paragraph.style.name == 'CET Address' and '@' not in paragraph.text:
                if self._is_contain_superscripts:
                    is_contain_superscripts = False
                    for r, run in enumerate(paragraph.runs):
                        if run.font.superscript is True:
                            is_contain_superscripts = True
                            superscript = run.text.strip()
                            break
                    
                    if is_contain_superscripts:
                        for text in paragraph.text[1:].strip().split("\n"):
                            affiliation_to_used = text
                            affiliations[superscript] = affiliation_to_used if superscript not in affiliations else affiliations[superscript] + affiliation_to_used 
                    else:
                        for t, text in enumerate(paragraph.text.strip().split("\n")):
                            affiliation_to_used = text
                            affiliations[superscript] += text if (t == 0) else text

                else:
                    for text in paragraph.text.strip().split("\n"):
                        if affiliations:
                            affiliations = list(affiliations)
                            affiliations[-1] += text
                            affiliations = set(affiliations)
                        else:
                            affiliations = {text}
        
        print(f'Affiliations: {affiliations}')
        return affiliations

    def _get_email(self, paragraphs):
        for p, paragraph in enumerate(paragraphs):
            if (

                ('@' in paragraph.text) and
                paragraph.text != ' '
            ):
                print(f'Email: {paragraph.text.strip()}')
                return paragraph.text.strip()
        pass

    def _get_authors_names_2(self, paragraph):
        if '*' not in paragraph.text:
            raise Exception('No corresponding authors!')
        authors = []
        corresponding_author = paragraph.text.strip().split('*')[0]
        corresponding_author = corresponding_author.split(',')
        author_label = []
        for r in range(1, len(corresponding_author) + 1):
            if len(corresponding_author[-r]) > 1:
                corresponding_author = corresponding_author[-r]
                break
            elif len(corresponding_author[-r]) == 1 and ' ' != corresponding_author[-r]:
                author_label.append(corresponding_author[-r])

        
        authors = [word.strip().split('*')[0] for word in re.split(',|，', paragraph.text) if len(word.strip()) > 1] # Split the text, and avoid getting the superscripts, and split from '*'

        is_contain_superscripts = False
        for r, run in enumerate(paragraph.runs):
            if run.font.superscript is True and run.text != '*':
                is_contain_superscripts = True
                break
        
        if is_contain_superscripts:
            authors = [author[:-1] for author in authors]
            author_label.append(corresponding_author[-1])
            corresponding_author = [corresponding_author[:-1]] 
            pass
        
        print(authors)
        self._is_contain_superscripts = is_contain_superscripts
        corresponding_author = [corresponding_author] if type(corresponding_author) != list else corresponding_author
        return Authors(author_list = authors, corresponding_author = corresponding_author), author_label


    def _get_authors_names(self, paragraph):
    # ---------------------------------------
    # Author names
    #----------------------------------------
        authors = []
        new_author = True
        
        for r, run in enumerate(paragraph.runs):
            # authors = [word.strip() for word in paragraph.text.split(',')]
            if (
                (run.font.superscript is None and '*' != run.text.strip()) and # Avoid superscript and only '*'
                '' != run.text.strip() or # If is not empty texts
                ('’' == run.text.strip() or '.' == run.text.strip() or '-' == run.text.strip()) # If only the special characters in authors name, can be included
            ): 
                if ("," in run.text):
                    check = run.text.split(',')
                    for text in run.text.split(','):
                        if text and len(text) > 1:
                            if ( 
                                not new_author
                            ):
                                authors[-1] += ' ' + text.strip().split('*')[0] if ('’' != run.text.strip() and '.' != run.text.strip() and '-' != run.text.strip()) else text.strip().split('*')[0]
                            else:
                                authors += [text.strip().split('*')[0].strip()]
                                new_author = False
                        
                        # Condition, if run.text is [', xxx' or 'yyy , xxx'] or just ','
                        new_author = True if (text != run.text.split(',')[-1] or '' == run.text.split(',')[-1]) else new_author

                else:
                    if ( 
                        not new_author
                    ):
                        authors[-1] += ' ' + run.text.strip().split('*')[0] if ('’' != run.text.strip() and '.' != run.text.strip() and '-' != run.text.strip()) else run.text.strip().split('*')[0]
                    else:
                        authors += [run.text.strip().split('*')[0].strip()]
                        new_author = False
            
            elif '' == run.text.strip():  # Ignore empty texts
                continue
            else:
                new_author = True

        print(authors)       
        return Authors(author_list = authors)
    
class CETManuscripts():
    all_info: List[CETExtraction]

    def __init__(self, file_list: List[str], file_path: str):
        self.all_info = []
        for file_name in file_list:
            self.all_info.append(CETExtraction(filename = file_name)) 
    
    def write_to_excel(self, file_path: str):
        rows_of_data_in_excel_lavori = []
        rows_of_data_in_excel_corresponding = []
        for manuscript_info in self.all_info:
        #______________________
        #LAVORI
        #______________________
            # (1): paper title
            # (2): page count
            # (3): paper id.pdf
            # (4): number of authors
            # (5): first name of first author
            # (6): last name of first author

            row = [
                manuscript_info.manuscript_info.paper_title,
                manuscript_info.manuscript_info.page_no,
                manuscript_info.paper_id,
                manuscript_info.authors.no_of_authors
            ]
            for author_ind in range(0, manuscript_info.authors.no_of_authors):
                row.append(manuscript_info.authors.first_name[author_ind])
                row.append(manuscript_info.authors.last_name[author_ind])

            rows_of_data_in_excel_lavori.append(row)
        
        #______________________
        #CORRESPONDING
        #______________________        
            # (1): paper title
            # (2): first name
            # (3): last name
            # (4): affiliation
            # (5): email
            row2 = [
                manuscript_info.manuscript_info.paper_title,
                manuscript_info.authors.corresponding_author.first_name[0],
                manuscript_info.authors.corresponding_author.last_name[0],
                manuscript_info.authors.corresponding_author.affiliation,
                manuscript_info.authors.corresponding_author.email,
            ]

            rows_of_data_in_excel_corresponding.append(row2)
        
        os.makedirs('downloads')

 
        df = pd.DataFrame(rows_of_data_in_excel_lavori)
        # df.to_excel(f"downloads//PRES23_CET_Info.xlsx", sheet_name = 'LAVORI')
       
        df2 = pd.DataFrame(rows_of_data_in_excel_corresponding)
        with pd.ExcelWriter(f"downloads//PRES23_CET_Info.xlsx", engine='xlsxwriter') as writer:
            df.to_excel(writer, sheet_name = 'LAVORI')
            df2.to_excel(writer, sheet_name = 'CORRESPONDING')
            # workbook  = writer.book
            # worksheet = writer.sheets['CORRESPONDING']
            # cell_format = workbook.add_format({'text_wrap': True})
            # worksheet.set_column('A:Z', cell_format=cell_format)
        # df2.to_excel(f"downloads//PRES23_CET_Info.xlsx", sheet_name = 'CORRESPONDING')
        
        pass

# Creating a Web App
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'

# @app.route('/', methods = ['GET'])
def get_CET_info(path: str = None, request = None):
    try:
        # Run on local
        # file_list = [file for file in os.listdir(path) if ('docx' in file and '$' not in file and 'PRES23' in file)]

        file_list = request.files.getlist('files[]')

        CET_manuscripts = CETManuscripts(file_list = file_list, file_path = path)
        CET_manuscripts.write_to_excel(path)

        response = {
            'message': f"Success! All CET info in the folder is extracted and saved to {path}/PRES23_CET_Info.xlsx"
        }

        return jsonify(response), 200

    except Exception as e:
        response = {
            'message': f"Errors! {e}"
        }
        return str(e), 400

@app.route('/', methods=['GET', 'POST'])
def get_folder_path():
    if request.method == 'POST':
        # Remove the folders before it is started
        shutil.rmtree('downloads', ignore_errors= True)

        response = get_CET_info(path = None, request = request)
        if response[1] == 200:
            return render_template('index.html', success = True)
        else:
            return render_template('index.html', error = True, message = response[0])
    return render_template('index.html')

@app.route('/downloads/<filename>')
def download_file(filename):
    
    return send_from_directory('downloads', filename, as_attachment=True)

if __name__ == "__main__":
    app.run(host = '0.0.0.0', debug = False)
    # path = askdirectory(title='Select Folder') # shows dialog box and return the path
    # # file_list = [file for file in os.listdir(path) if ('docx' in file and '$' not in file)]
    # file_list = [file for file in os.listdir(path) if ('docx' in file and '$' not in file and 'PRES23' in file)]
    # # file_list = ["+PRES23_0184_M_v_03_Author.docx"]
    # # file_list = ["+PRES23_0340_rev_M_v2_Review.docx"]
    # # file_list = ["+PRES23_0002_rev_M_v2_Review.docx"]
    # CET_manuscripts = CETManuscripts(file_list = file_list, file_path = path)
    # CET_manuscripts.write_to_excel(path)

    pass
